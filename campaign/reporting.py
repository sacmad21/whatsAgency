from docx import Document
from docx.shared import Pt
from datetime import datetime



from util import get_research_and_objective_setting_data, format_date

def get_stage_report_filename(stage_name: str, campaign_id: str) -> str:
    """
    Constructs a consistent file path for a stage report based on stage name and campaign ID.
    Output files are stored in subdirectory: ./reports/<stage_name>/
    """
    import os

    safe_stage = stage_name.strip().lower().replace(" ", "_")
    directory = os.path.join("reports/campaign", safe_stage)

    os.makedirs(directory, exist_ok=True)

    filename = f"{safe_stage}_{campaign_id}.docx"
    return os.path.join(directory, filename)


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_table(doc: Document, data: dict, title: str = None):
    if title:
        doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'

    for key, val in data.items():
        row = table.add_row().cells
        row[0].text = key.replace("_", " ").title()
        row[1].text = str(val)
    doc.add_paragraph("")


def clean_dict_for_report(data: dict) -> dict:
    """
    Removes irrelevant or repetitive keys for better document readability.
    """
    ignore_keys = {"id", "campaignid", "campaign", "createdat", "created_date"}
    return {
        key.replace("_", " ").title(): value
        for key, value in data.items()
        if value not in [None, "", "None"] and key.lower() not in ignore_keys
    }


def add_paragraph_block(doc: Document, title: str, items: list[dict]):
    if not items:
        return
    doc.add_heading(title, level=2)
    for i, item in enumerate(items, 1):
        cleaned = clean_dict_for_report(item.dict())
        lines = [f"{k}: {v}" for k, v in cleaned.items()]
        doc.add_paragraph(f"{i}. " + "; ".join(lines))


async def export_research_and_objective_setting_report(campaign_id: str) -> str:
    """
    Export 'Research and Objective Setting' stage data as a clean, narrative Word document.
    """
    data = await get_research_and_objective_setting_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Research and Objective Setting", 0)
    doc.add_paragraph("Goal: Define the purpose and scope of the campaign.", style="Intense Quote")

    add_paragraph_block(doc, "Campaign Objectives", data["campaignobjective"])
    add_paragraph_block(doc, "Target Audience Segments", data["audiencesegment"])
    add_paragraph_block(doc, "Competitor Strategies", data["competitorstrategy"])
    add_paragraph_block(doc, "Promotion Types", data["promotiontype"])

    doc.add_paragraph(f"\nReport generated on: {format_date(datetime.now())}", style="Normal")

    filename = get_stage_report_filename("Research and Objective Setting", campaign_id)
    doc.save(filename)
    return filename




async def export_research_and_objective_setting_report_V2(campaign_id: str) -> str:
    """
    Export 'Research and Objective Setting' stage data to a Word document.
    """
    data = await get_research_and_objective_setting_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Research and Objective Setting", 0)
    doc.add_paragraph("Goal: Define the purpose and scope of the campaign.", style="Intense Quote")

    add_heading(doc, "Campaign Objectives")
    for i, item in enumerate(data["campaignobjective"], 1):
        add_table(doc, item.dict(), title=f"Objective {i}")

    add_heading(doc, "Target Audience Segments")
    for i, item in enumerate(data["audiencesegment"], 1):
        add_table(doc, item.dict(), title=f"Segment {i}")

    add_heading(doc, "Competitor Strategies")
    for i, item in enumerate(data["competitorstrategy"], 1):
        add_table(doc, item.dict(), title=f"Competitor {i}")

    add_heading(doc, "Promotion Types")
    for i, item in enumerate(data["promotiontype"], 1):
        add_table(doc, item.dict(), title=f"Promotion {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    filename = get_stage_report_filename("research_objective_stage", campaign_id)
    doc.save(filename)
    return filename




from util import get_campaign_planning_data


async def export_campaign_planning_report(campaign_id: str) -> str:
    """
    Export 'Campaign Planning' stage data to a Word document.
    """
    data = await get_campaign_planning_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Campaign Planning", 0)
    doc.add_paragraph("Goal: Design the mechanics and structure of the campaign.", style="Intense Quote")

    add_heading(doc, "Campaign Offers")
    for i, item in enumerate(data["campaignoffer"], 1):
        add_table(doc, item.dict(), title=f"Offer {i}")

    add_heading(doc, "Campaign Budgets")
    for i, item in enumerate(data["campaignbudget"], 1):
        add_table(doc, item.dict(), title=f"Budget {i}")

    add_heading(doc, "Channel Plans")
    for i, item in enumerate(data["channelplan"], 1):
        add_table(doc, item.dict(), title=f"Channel Plan {i}")

    add_heading(doc, "Campaign Timelines")
    for i, item in enumerate(data["campaigntimeline"], 1):
        add_table(doc, item.dict(), title=f"Timeline {i}")

    add_heading(doc, "Compliance Checklist")
    for i, item in enumerate(data["compliancechecklist"], 1):
        add_table(doc, item.dict(), title=f"Compliance Item {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    filename = get_stage_report_filename("campaign_planning_stage", campaign_id)


    doc.save(filename)
    return filename











from util import get_content_creation_and_design_data


async def export_content_creation_and_design_report(campaign_id: str) -> str:
    """
    Export 'Content Creation and Design' stage data to a Word document.
    """
    data = await get_content_creation_and_design_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Content Creation and Design", 0)
    doc.add_paragraph("Goal: Develop compelling promotional material to attract customers.", style="Intense Quote")

    add_heading(doc, "Promotional Messages")
    for i, item in enumerate(data["promotionalmessage"], 1):
        add_table(doc, item.dict(), title=f"Message {i}")

    add_heading(doc, "Creative Assets")
    for i, item in enumerate(data["creativeasset"], 1):
        add_table(doc, item.dict(), title=f"Creative {i}")

    add_heading(doc, "Media Assets")
    for i, item in enumerate(data["mediaasset"], 1):
        add_table(doc, item.dict(), title=f"Media {i}")

    add_heading(doc, "Content Calendar")
    for i, item in enumerate(data["contentcalendar"], 1):
        add_table(doc, item.dict(), title=f"Calendar Entry {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    filename = get_stage_report_filename("content_creation_stage", campaign_id)

    doc.save(filename)
    return filename















from util import get_pre_launch_phase_data


async def export_pre_launch_phase_report(campaign_id: str) -> str:
    """
    Export 'Pre-Launch Phase' stage data to a Word document.
    """
    data = await get_pre_launch_phase_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Pre-Launch Phase", 0)
    doc.add_paragraph("Goal: Build anticipation and ensure operational readiness for the campaign.", style="Intense Quote")

    add_heading(doc, "Teaser Content")
    for i, item in enumerate(data["teasercontent"], 1):
        add_table(doc, item.dict(), title=f"Teaser {i}")

    add_heading(doc, "Customer Segments")
    for i, item in enumerate(data["customersegmentlist"], 1):
        add_table(doc, item.dict(), title=f"Segment {i}")

    add_heading(doc, "Influencer Plans")
    for i, item in enumerate(data["influencerplan"], 1):
        add_table(doc, item.dict(), title=f"Influencer {i}")

    add_heading(doc, "Operational Checklists")
    for i, item in enumerate(data["operationalchecklist"], 1):
        add_table(doc, item.dict(), title=f"Checklist Item {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


    filename = get_stage_report_filename("pre_launch_phase_stage", campaign_id)

    doc.save(filename)
    return filename













from util import get_campaign_launch_data


async def export_campaign_launch_report(campaign_id: str) -> str:
    """
    Export 'Campaign Launch' stage data to a Word document.
    """
    data = await get_campaign_launch_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Campaign Launch", 0)
    doc.add_paragraph("Goal: Execute the campaign and drive customer action.", style="Intense Quote")

    add_heading(doc, "Campaign Activations")
    for i, item in enumerate(data["campaignactivation"], 1):
        add_table(doc, item.dict(), title=f"Activation {i}")

    add_heading(doc, "Customer Engagement Logs")
    for i, item in enumerate(data["customerengagementlog"], 1):
        add_table(doc, item.dict(), title=f"Engagement {i}")

    add_heading(doc, "Performance Reports")
    for i, item in enumerate(data["performancereport"], 1):
        add_table(doc, item.dict(), title=f"Report {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


    filename = get_stage_report_filename("campaign_launch_stage", campaign_id)

    doc.save(filename)
    return filename












from util import get_post_launch_follow_up_data


async def export_post_launch_follow_up_report(campaign_id: str) -> str:
    """
    Export 'Post-Launch Follow-Up' stage data to a Word document.
    """
    data = await get_post_launch_follow_up_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Post-Launch Follow-Up", 0)
    doc.add_paragraph("Goal: Sustain interest and convert leads into loyal customers.", style="Intense Quote")

    add_heading(doc, "Customer Feedback")
    for i, item in enumerate(data["customerfeedback"], 1):
        add_table(doc, item.dict(), title=f"Feedback {i}")

    add_heading(doc, "Thank You Messages")
    for i, item in enumerate(data["thankyoumessage"], 1):
        add_table(doc, item.dict(), title=f"Thank You {i}")

    add_heading(doc, "Retargeting Plans")
    for i, item in enumerate(data["retargetingplan"], 1):
        add_table(doc, item.dict(), title=f"Retargeting {i}")

    add_heading(doc, "Campaign Extensions")
    for i, item in enumerate(data["campaignextensionplan"], 1):
        add_table(doc, item.dict(), title=f"Extension {i}")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


    filename = get_stage_report_filename("post_launch_follow_up_stage", campaign_id)

    doc.save(filename)
    return filename














from util import get_campaign_analysis_data


async def export_campaign_analysis_report(campaign_id: str) -> str:
    """
    Export 'Campaign Analysis' stage data to a Word document.
    """
    data = await get_campaign_analysis_data(campaign_id)
    doc = Document()

    doc.add_heading("Stage: Campaign Analysis", 0)
    doc.add_paragraph("Goal: Evaluate the success of the campaign and identify areas for improvement.", style="Intense Quote")

    add_heading(doc, "Campaign Analysis Report")
    if data["campaignanalysisreport"]:
        add_table(doc, data["campaignanalysisreport"].dict(), title="Campaign Analysis")

    add_heading(doc, "Campaign Learnings")
    if data["campaignlearnings"]:
        add_table(doc, data["campaignlearnings"].dict(), title="Key Learnings")

    add_heading(doc, "Internal Campaign Report")
    if data["internalcampaignreport"]:
        add_table(doc, data["internalcampaignreport"].dict(), title="Internal Report")

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


    filename = get_stage_report_filename("campaign_analysis_stage", campaign_id)

    doc.save(filename)
    return filename










async def export_stage_report_by_name(stage_name: str, campaign_id: str) -> str:
    """
    Exports a specific stage report by name for a given campaign_id.
    Returns the path of the generated .docx file.
    """
    stage_map = {
        "research_and_objective_setting": export_research_and_objective_setting_report,
        "campaign_planning": export_campaign_planning_report,
        "content_creation_and_design": export_content_creation_and_design_report,
        "pre_launch_phase": export_pre_launch_phase_report,
        "campaign_launch": export_campaign_launch_report,
        "post_launch_follow_up": export_post_launch_follow_up_report,
        "campaign_analysis": export_campaign_analysis_report,
    }

    normalized = stage_name.strip().lower().replace(" ", "_")
    export_func = stage_map.get(normalized)

    if not export_func:
        raise ValueError(f"Unknown stage name: {stage_name}")

    return await export_func(campaign_id)






































import asyncio


async def export_all_campaign_reports(campaign_id: str) -> list:
    """
    Runs all stage-specific export functions for a given campaign_id.
    Returns list of generated file paths.
    """
    reports = []

    reports.append(await export_research_and_objective_setting_report(campaign_id))
    reports.append(await export_campaign_planning_report(campaign_id))
    reports.append(await export_content_creation_and_design_report(campaign_id))
    reports.append(await export_pre_launch_phase_report(campaign_id))
    reports.append(await export_campaign_launch_report(campaign_id))
    reports.append(await export_post_launch_follow_up_report(campaign_id))
    reports.append(await export_campaign_analysis_report(campaign_id))

    return reports

